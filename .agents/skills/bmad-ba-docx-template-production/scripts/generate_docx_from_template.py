import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


PLACEHOLDER = re.compile(r"{{\s*([a-zA-Z0-9_.-]+)\s*}}")


def get_value(data, path, default="[Chưa cung cấp]"):
    current = data
    for part in path.split("."):
        if isinstance(current, dict) and part in current:
            current = current[part]
        else:
            return default
    if current is None or current == "":
        return default
    return current


def render_text(text, data, row=None):
    context = data.copy()
    if row:
        context.update(row)
    return PLACEHOLDER.sub(lambda match: str(get_value(context, match.group(1))), text or "")


def alignment(value, paragraph=False):
    paragraph_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    table_map = {
        "left": WD_TABLE_ALIGNMENT.LEFT,
        "center": WD_TABLE_ALIGNMENT.CENTER,
        "right": WD_TABLE_ALIGNMENT.RIGHT,
    }
    return (paragraph_map if paragraph else table_map).get(value or "left")


def set_cell_shading(cell, fill):
    if not fill:
        return
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="808080"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:%s" % edge
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:color"), color or "808080")
        element.set(qn("w:space"), "0")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, font_name, style):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.bold = bool(style.get("bold", False))
    run.italic = bool(style.get("italic", False))
    if style.get("fontSize"):
        run.font.size = Pt(float(style["fontSize"]))
    if style.get("color"):
        run.font.color.rgb = RGBColor.from_string(style["color"])


def apply_paragraph_style(paragraph, style):
    if style.get("alignment"):
        paragraph.alignment = alignment(style["alignment"], paragraph=True)
    fmt = paragraph.paragraph_format
    if style.get("spacingBefore") is not None:
        fmt.space_before = Pt(float(style["spacingBefore"]))
    if style.get("spacingAfter") is not None:
        fmt.space_after = Pt(float(style["spacingAfter"]))
    if style.get("lineSpacing") is not None:
        fmt.line_spacing = float(style["lineSpacing"])


def style_for(template, kind, name):
    return template.get("styles", {}).get(kind, {}).get(name or "", {})


def add_paragraph(doc, template, data, block):
    page = template.get("page", {})
    style = style_for(template, "paragraphStyles", block.get("style"))
    texts = get_value(data, block["source"], []) if block.get("source") and not block.get("text") else [render_text(block.get("text", ""), data)]
    if not isinstance(texts, list):
        texts = [texts]
    for text in texts:
        paragraph = doc.add_paragraph()
        apply_paragraph_style(paragraph, style)
        run = paragraph.add_run(str(text))
        set_font(run, page.get("defaultFont", "Times New Roman"), {**{"fontSize": page.get("defaultFontSize", 12)}, **style})


def add_heading(doc, template, data, block):
    paragraph = doc.add_paragraph()
    style = style_for(template, "paragraphStyles", block.get("style"))
    apply_paragraph_style(paragraph, style)
    run = paragraph.add_run(render_text(block.get("text", ""), data))
    page = template.get("page", {})
    set_font(run, page.get("defaultFont", "Times New Roman"), {**{"bold": True, "fontSize": 13}, **style})


def add_bullet_list(doc, template, data, block):
    page = template.get("page", {})
    style = style_for(template, "paragraphStyles", block.get("style"))
    items = get_value(data, block.get("source", ""), [])
    if not isinstance(items, list):
        items = [items]
    for item in items:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.6)
        paragraph.paragraph_format.first_line_indent = Cm(-0.3)
        apply_paragraph_style(paragraph, style)
        run = paragraph.add_run("- " + str(item))
        set_font(run, page.get("defaultFont", "Times New Roman"), {**{"fontSize": page.get("defaultFontSize", 12)}, **style})


def normalize_rows(value, columns):
    if isinstance(value, dict):
        return [{"label": key, "value": val} for key, val in value.items()]
    if isinstance(value, list):
        rows = []
        for index, item in enumerate(value, start=1):
            row = item.copy() if isinstance(item, dict) else {"value": item}
            row.setdefault("index", index)
            rows.append(row)
        return rows
    if value in (None, "[Chưa cung cấp]"):
        return []
    first_key = columns[0].get("key", "value") if columns else "value"
    return [{first_key: value}]


def add_table(doc, template, data, block):
    columns = block.get("columns", [])
    rows = normalize_rows(get_value(data, block.get("source", ""), []), columns)
    table_style = style_for(template, "tableStyles", block.get("style"))
    include_header = not (block.get("source") and isinstance(get_value(data, block.get("source"), None), dict))
    table = doc.add_table(rows=1 if include_header else 0, cols=max(len(columns), 1))
    table.alignment = alignment(table_style.get("alignment", "center"))
    table.style = "Table Grid"

    if include_header:
        for index, column in enumerate(columns):
            cell = table.cell(0, index)
            cell.text = column.get("header", "")
            set_cell_shading(cell, column.get("fill") or table_style.get("headerFill"))
            for paragraph in cell.paragraphs:
                paragraph.alignment = alignment(column.get("alignment", "center"), paragraph=True)
                for run in paragraph.runs:
                    set_font(run, template.get("page", {}).get("defaultFont", "Times New Roman"), {"bold": True, "fontSize": 11})
        if table_style.get("repeatHeader"):
            set_repeat_table_header(table.rows[0])

    for row_data in rows:
        row = table.add_row()
        for index, column in enumerate(columns):
            cell = row.cells[index]
            cell.width = Cm(float(column.get("widthCm", 4)))
            cell.text = str(row_data.get(column.get("key"), ""))
            set_cell_shading(cell, column.get("fill"))
            for paragraph in cell.paragraphs:
                if column.get("alignment"):
                    paragraph.alignment = alignment(column.get("alignment"), paragraph=True)

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell, table_style.get("borderColor", "808080"))
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(3)
                paragraph.paragraph_format.line_spacing = 1.1
                for run in paragraph.runs:
                    set_font(run, template.get("page", {}).get("defaultFont", "Times New Roman"), {"fontSize": 11})


def add_signature(doc, template, data, block):
    table = doc.add_table(rows=2, cols=2)
    labels = [render_text(block.get("leftLabel", ""), data), render_text(block.get("rightLabel", ""), data)]
    table_style = style_for(template, "tableStyles", block.get("style"))
    table.alignment = alignment(table_style.get("alignment", "center"))
    table.style = "Table Grid"
    for index, label in enumerate(labels):
        header_cell = table.cell(0, index)
        sign_cell = table.cell(1, index)
        header_cell.text = label
        sign_cell.text = "\n\n\n(Ký, ghi rõ họ tên)"
        set_cell_shading(header_cell, table_style.get("headerFill", "EDEDED"))
        for paragraph in header_cell.paragraphs + sign_cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, table_style.get("borderColor", "808080"))


def setup_page(doc, template):
    page = template.get("page", {})
    section = doc.sections[0]
    if page.get("orientation") == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    margins = page.get("marginsCm", {})
    section.top_margin = Cm(float(margins.get("top", 2.0)))
    section.bottom_margin = Cm(float(margins.get("bottom", 2.0)))
    section.left_margin = Cm(float(margins.get("left", 2.5)))
    section.right_margin = Cm(float(margins.get("right", 2.0)))
    doc.styles["Normal"].font.name = page.get("defaultFont", "Times New Roman")
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), page.get("defaultFont", "Times New Roman"))
    doc.styles["Normal"].font.size = Pt(float(page.get("defaultFontSize", 12)))


def build_docx(template, data, output_path):
    doc = Document()
    setup_page(doc, template)
    handlers = {
        "paragraph": add_paragraph,
        "heading": add_heading,
        "bullet_list": add_bullet_list,
        "table": add_table,
        "signature": add_signature,
    }
    for block in template.get("blocks", []):
        block_type = block.get("type")
        if block_type == "spacer":
            doc.add_paragraph("")
        elif block_type == "page_break":
            doc.add_page_break()
        elif block_type == "image" and block.get("path"):
            doc.add_picture(render_text(block["path"], data), width=Cm(float(block.get("widthCm", 12))))
        elif block_type in handlers:
            handlers[block_type](doc, template, data, block)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser(description="Generate a DOCX file from a JSON template and JSON content.")
    parser.add_argument("template", help="Template JSON path")
    parser.add_argument("content", help="Content JSON path")
    parser.add_argument("--output", required=True, help="Output .docx path")
    args = parser.parse_args()

    template = json.loads(Path(args.template).read_text(encoding="utf-8"))
    data = json.loads(Path(args.content).read_text(encoding="utf-8"))
    output_path = Path(args.output)
    build_docx(template, data, output_path)
    print(output_path)


if __name__ == "__main__":
    main()
