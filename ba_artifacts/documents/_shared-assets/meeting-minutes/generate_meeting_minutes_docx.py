import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


FONT_NAME = "Times New Roman"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:color"), "808080")
        element.set(qn("w:space"), "0")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, bold=False, italic=False, size=12, color=None):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_text(doc, text="", bold=False, italic=False, size=12, align=None, spacing_after=6):
    paragraph = doc.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(spacing_after)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_font(run, bold=bold, italic=italic, size=size)
    return paragraph


def add_section_heading(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_font(run, bold=True, size=13)


def add_bullet(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.6)
    paragraph.paragraph_format.first_line_indent = Cm(-0.3)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run("- " + text)
    set_font(run, size=12)


def format_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(3)
                paragraph.paragraph_format.line_spacing = 1.1
                for run in paragraph.runs:
                    set_font(run, size=11)


def add_info_table(doc, meeting_info):
    table = doc.add_table(rows=len(meeting_info), cols=2)
    for index, (label, value) in enumerate(meeting_info.items()):
        label_cell = table.cell(index, 0)
        value_cell = table.cell(index, 1)
        label_cell.text = label
        value_cell.text = value
        set_cell_shading(label_cell, "EDEDED")
        for run in label_cell.paragraphs[0].runs:
            set_font(run, bold=True, size=11)
        for run in value_cell.paragraphs[0].runs:
            set_font(run, size=11)
    format_table(table)


def add_work_table(doc, items):
    headers = ["STT", "Nội dung công việc", "Đơn vị phụ trách", "Thông tin cần phản hồi/Thời hạn"]
    table = doc.add_table(rows=1, cols=4)
    widths = [Cm(1.2), Cm(8.3), Cm(2.7), Cm(4.3)]

    for index, header in enumerate(headers):
        cell = table.cell(0, index)
        cell.text = header
        cell.width = widths[index]
        set_cell_shading(cell, "D9EAF7")
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_font(run, bold=True, size=11)
    set_repeat_table_header(table.rows[0])

    for row_number, item in enumerate(items, start=1):
        row = table.add_row()
        values = [
            str(row_number),
            item.get("content", ""),
            item.get("owner", ""),
            item.get("deadline_or_response", ""),
        ]
        for index, value in enumerate(values):
            cell = row.cells[index]
            cell.text = value
            cell.width = widths[index]
            for paragraph in cell.paragraphs:
                if index in (0, 2):
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_font(run, size=11)
    format_table(table)


def add_signature_table(doc, left_label, right_label):
    table = doc.add_table(rows=2, cols=2)
    for index, label in enumerate([left_label, right_label]):
        header_cell = table.cell(0, index)
        sign_cell = table.cell(1, index)
        header_cell.text = label
        sign_cell.text = "\n\n\n(Ký, ghi rõ họ tên)"
        set_cell_shading(header_cell, "EDEDED")
        for paragraph in header_cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_font(run, bold=True, size=11)
        for paragraph in sign_cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_font(run, italic=True, size=11)
    format_table(table)


def build_docx(data, output_path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

    doc.styles["Normal"].font.name = FONT_NAME
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    doc.styles["Normal"].font.size = Pt(12)

    add_text(doc, data.get("title", "BIÊN BẢN NỘI DUNG CUỘC HỌP"), bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=3)
    add_text(doc, data.get("subtitle", ""), bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=12)

    add_info_table(doc, data.get("meeting_info", {}))

    add_section_heading(doc, "I. NỘI DUNG ĐÃ THỐNG NHẤT")
    for item in data.get("agreed_items", []):
        add_bullet(doc, item)

    add_section_heading(doc, "II. CÁC CÔNG VIỆC HỖ TRỢ XỬ LÝ TIẾP THEO")
    add_work_table(doc, data.get("follow_up_items", []))

    add_section_heading(doc, "III. KẾT LUẬN")
    for item in data.get("conclusion", []):
        add_text(doc, item, spacing_after=8)

    add_signature_table(
        doc,
        data.get("signature_left", "ĐẠI DIỆN KHÁCH HÀNG"),
        data.get("signature_right", "ĐẠI DIỆN FAST"),
    )
    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser(description="Generate a formal Vietnamese meeting minutes DOCX file.")
    parser.add_argument("input", help="Path to meeting minutes JSON input file.")
    parser.add_argument("--output", help="Output DOCX path. Overrides output_file in JSON.")
    args = parser.parse_args()

    input_path = Path(args.input)
    with input_path.open("r", encoding="utf-8") as file:
        data = json.load(file)

    if args.output:
        output_path = Path(args.output)
    else:
        output_path = Path(data.get("output_file", "Bien_ban_noi_dung_cuoc_hop.docx"))
    if not args.output and not output_path.is_absolute():
        output_path = input_path.parent / output_path

    build_docx(data, output_path)
    print(output_path)


if __name__ == "__main__":
    main()
